"""
Document generation for Career Manager.
Generates PDF, DOCX, and Markdown resume files.

Dependencies (optional):
- python-docx: for DOCX generation
- reportlab: for PDF generation

If dependencies are not installed, only Markdown generation is available.
"""

import os
from pathlib import Path
from datetime import datetime
from typing import Literal
import json

from .storage import load_data

# Default output directory
OUTPUT_DIR = Path.home() / ".career-manager" / "output"


def ensure_output_dir() -> Path:
    """Ensure output directory exists and return path."""
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    return OUTPUT_DIR


def generate_resume_content() -> dict:
    """Generate resume content from stored data.

    Returns dict with all resume sections ready for document generation.
    """
    profile = load_data("profile")
    resume = load_data("resume")

    personal = profile.get("personal_info", {})
    social = profile.get("social_links", {})
    sections = resume.get("sections", {})

    return {
        "header": {
            "name": personal.get("name", "Your Name"),
            "title": personal.get("title", ""),
            "email": personal.get("email", ""),
            "phone": personal.get("phone", ""),
            "location": personal.get("location", ""),
            "linkedin": social.get("linkedin", ""),
            "github": social.get("github", ""),
            "portfolio": social.get("portfolio", ""),
            "website": social.get("website", ""),
        },
        "summary": sections.get("summary", ""),
        "experience": sections.get("experience", []),
        "education": sections.get("education", []),
        "skills": sections.get("skills", []),
        "certifications": sections.get("certifications", []),
        "projects": sections.get("projects", []),
    }


def format_date(date_str: str | None) -> str:
    """Format a date string for display."""
    if not date_str:
        return ""
    if date_str.lower() == "present":
        return "Present"
    try:
        # Handle YYYY-MM format
        if len(date_str) == 7:
            year, month = date_str.split("-")
            months = [
                "Jan", "Feb", "Mar", "Apr", "May", "Jun",
                "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"
            ]
            return f"{months[int(month) - 1]} {year}"
        return date_str
    except (ValueError, IndexError):
        return date_str


# ============== MARKDOWN GENERATION ==============

def generate_markdown(content: dict = None, filename: str = None) -> Path:
    """Generate a Markdown resume file.

    Args:
        content: Resume content dict (generated if not provided)
        filename: Output filename (auto-generated if not provided)

    Returns:
        Path to the generated file
    """
    if content is None:
        content = generate_resume_content()

    ensure_output_dir()

    if filename is None:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"resume_{timestamp}.md"

    filepath = OUTPUT_DIR / filename

    lines = []

    # Header
    header = content["header"]
    lines.append(f"# {header['name']}")
    if header.get("title"):
        lines.append(f"**{header['title']}**")
    lines.append("")

    # Contact info
    contacts = []
    if header.get("email"):
        contacts.append(f"📧 {header['email']}")
    if header.get("phone"):
        contacts.append(f"📱 {header['phone']}")
    if header.get("location"):
        contacts.append(f"📍 {header['location']}")

    if contacts:
        lines.append(" | ".join(contacts))
        lines.append("")

    # Social links
    socials = []
    if header.get("linkedin"):
        socials.append(f"[LinkedIn]({header['linkedin']})")
    if header.get("github"):
        socials.append(f"[GitHub]({header['github']})")
    if header.get("portfolio"):
        socials.append(f"[Portfolio]({header['portfolio']})")
    if header.get("website"):
        socials.append(f"[Website]({header['website']})")

    if socials:
        lines.append(" | ".join(socials))
        lines.append("")

    # Summary
    if content.get("summary"):
        lines.append("## Summary")
        lines.append("")
        lines.append(content["summary"])
        lines.append("")

    # Experience
    if content.get("experience"):
        lines.append("## Experience")
        lines.append("")
        for exp in content["experience"]:
            title = exp.get("title", "")
            company = exp.get("company", "")
            lines.append(f"### {title} at {company}")

            dates = f"{format_date(exp.get('start_date'))} - {format_date(exp.get('end_date'))}"
            location = exp.get("location", "")
            if location:
                lines.append(f"*{dates} | {location}*")
            else:
                lines.append(f"*{dates}*")
            lines.append("")

            if exp.get("description"):
                lines.append(exp["description"])
                lines.append("")

            if exp.get("achievements"):
                lines.append("**Key Achievements:**")
                for achievement in exp["achievements"]:
                    lines.append(f"- {achievement}")
                lines.append("")

            if exp.get("technologies"):
                lines.append(f"*Technologies: {', '.join(exp['technologies'])}*")
                lines.append("")

    # Education
    if content.get("education"):
        lines.append("## Education")
        lines.append("")
        for edu in content["education"]:
            degree = edu.get("degree", "")
            field = edu.get("field", "")
            institution = edu.get("institution", "")

            if field:
                lines.append(f"### {degree} in {field}")
            else:
                lines.append(f"### {degree}")
            lines.append(f"{institution}")

            dates = f"{format_date(edu.get('start_date'))} - {format_date(edu.get('end_date'))}"
            if dates != " - ":
                lines.append(f"*{dates}*")

            if edu.get("gpa"):
                lines.append(f"GPA: {edu['gpa']}")

            if edu.get("achievements"):
                for achievement in edu["achievements"]:
                    lines.append(f"- {achievement}")
            lines.append("")

    # Skills
    if content.get("skills"):
        lines.append("## Skills")
        lines.append("")

        skills_by_cat = {}
        for skill in content["skills"]:
            cat = skill.get("category", "technical")
            if cat not in skills_by_cat:
                skills_by_cat[cat] = []
            skills_by_cat[cat].append(skill)

        for cat in ["technical", "soft", "language"]:
            if cat in skills_by_cat:
                cat_name = cat.title() + (" Skills" if cat != "language" else "s")
                skills_list = []
                for skill in skills_by_cat[cat]:
                    skill_str = skill.get("name", "")
                    if skill.get("level"):
                        skill_str += f" ({skill['level']})"
                    if skill.get("years"):
                        skill_str += f" - {skill['years']} years"
                    skills_list.append(skill_str)
                lines.append(f"**{cat_name}:** {', '.join(skills_list)}")
                lines.append("")

    # Certifications
    if content.get("certifications"):
        lines.append("## Certifications")
        lines.append("")
        for cert in content["certifications"]:
            name = cert.get("name", "")
            issuer = cert.get("issuer", "")
            date = format_date(cert.get("date"))

            line = f"- **{name}**"
            if issuer:
                line += f" - {issuer}"
            if date:
                line += f" ({date})"
            lines.append(line)
        lines.append("")

    # Projects
    if content.get("projects"):
        lines.append("## Projects")
        lines.append("")
        for proj in content["projects"]:
            name = proj.get("name", "")
            lines.append(f"### {name}")

            if proj.get("description"):
                lines.append(proj["description"])
                lines.append("")

            if proj.get("technologies"):
                lines.append(f"*Technologies: {', '.join(proj['technologies'])}*")
                lines.append("")

            links = []
            if proj.get("url"):
                links.append(f"[Live Demo]({proj['url']})")
            if proj.get("github"):
                links.append(f"[GitHub]({proj['github']})")
            if links:
                lines.append(" | ".join(links))
                lines.append("")

    # Write file
    filepath.write_text("\n".join(lines), encoding="utf-8")
    return filepath


# ============== DOCX GENERATION ==============

def generate_docx(content: dict = None, filename: str = None) -> Path:
    """Generate a DOCX resume file.

    Requires python-docx package.

    Args:
        content: Resume content dict (generated if not provided)
        filename: Output filename (auto-generated if not provided)

    Returns:
        Path to the generated file
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError(
            "python-docx is required for DOCX generation. "
            "Install with: pip install python-docx"
        )

    if content is None:
        content = generate_resume_content()

    ensure_output_dir()

    if filename is None:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"resume_{timestamp}.docx"

    filepath = OUTPUT_DIR / filename

    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Header
    header = content["header"]

    # Name
    name_para = doc.add_paragraph()
    name_run = name_para.add_run(header["name"])
    name_run.bold = True
    name_run.font.size = Pt(18)
    name_run.font.color.rgb = RGBColor(31, 73, 125)
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Title
    if header.get("title"):
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(header["title"])
        title_run.font.size = Pt(12)
        title_run.font.color.rgb = RGBColor(89, 89, 89)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Contact
    contact_parts = []
    if header.get("email"):
        contact_parts.append(header["email"])
    if header.get("phone"):
        contact_parts.append(header["phone"])
    if header.get("location"):
        contact_parts.append(header["location"])

    if contact_parts:
        contact_para = doc.add_paragraph()
        contact_para.add_run(" | ".join(contact_parts))
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Social links
    social_parts = []
    if header.get("linkedin"):
        social_parts.append(f"LinkedIn: {header['linkedin']}")
    if header.get("github"):
        social_parts.append(f"GitHub: {header['github']}")
    if header.get("portfolio"):
        social_parts.append(f"Portfolio: {header['portfolio']}")

    if social_parts:
        social_para = doc.add_paragraph()
        social_para.add_run(" | ".join(social_parts))
        social_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # Spacer

    # Helper function to add section heading
    def add_heading(text: str):
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(31, 73, 125)
        para.space_after = Pt(6)

    # Summary
    if content.get("summary"):
        add_heading("SUMMARY")
        summary_para = doc.add_paragraph()
        summary_para.add_run(content["summary"])
        doc.add_paragraph()

    # Experience
    if content.get("experience"):
        add_heading("EXPERIENCE")

        for exp in content["experience"]:
            # Company and title
            exp_title = doc.add_paragraph()
            title_run = exp_title.add_run(exp.get("title", ""))
            title_run.bold = True

            company_run = exp_title.add_run(f" | {exp.get('company', '')}")
            company_run.font.color.rgb = RGBColor(89, 89, 89)

            # Dates and location
            dates_para = doc.add_paragraph()
            dates = f"{format_date(exp.get('start_date'))} - {format_date(exp.get('end_date'))}"
            if exp.get("location"):
                dates += f" | {exp['location']}"
            dates_run = dates_para.add_run(dates)
            dates_run.italic = True
            dates_run.font.color.rgb = RGBColor(128, 128, 128)

            # Description
            if exp.get("description"):
                desc_para = doc.add_paragraph()
                desc_para.add_run(exp["description"])

            # Achievements
            if exp.get("achievements"):
                for achievement in exp["achievements"]:
                    bullet = doc.add_paragraph(style="List Bullet")
                    bullet.add_run(achievement)

            # Technologies
            if exp.get("technologies"):
                tech_para = doc.add_paragraph()
                tech_run = tech_para.add_run(f"Technologies: {', '.join(exp['technologies'])}")
                tech_run.italic = True

            doc.add_paragraph()  # Spacer

    # Education
    if content.get("education"):
        add_heading("EDUCATION")

        for edu in content["education"]:
            degree_para = doc.add_paragraph()
            degree_run = degree_para.add_run(edu.get("degree", ""))
            degree_run.bold = True

            if edu.get("field"):
                degree_para.add_run(f" in {edu['field']}")

            institution_para = doc.add_paragraph()
            institution_para.add_run(edu.get("institution", ""))

            if edu.get("gpa"):
                gpa_para = doc.add_paragraph()
                gpa_para.add_run(f"GPA: {edu['gpa']}")

            doc.add_paragraph()

    # Skills
    if content.get("skills"):
        add_heading("SKILLS")

        skills_by_cat = {}
        for skill in content["skills"]:
            cat = skill.get("category", "technical")
            if cat not in skills_by_cat:
                skills_by_cat[cat] = []
            skills_by_cat[cat].append(skill)

        for cat in ["technical", "soft", "language"]:
            if cat in skills_by_cat:
                cat_name = cat.title() + (" Skills" if cat != "language" else "s")
                skills_para = doc.add_paragraph()

                cat_run = skills_para.add_run(f"{cat_name}: ")
                cat_run.bold = True

                skill_names = [s.get("name", "") for s in skills_by_cat[cat]]
                skills_para.add_run(", ".join(skill_names))

        doc.add_paragraph()

    # Certifications
    if content.get("certifications"):
        add_heading("CERTIFICATIONS")

        for cert in content["certifications"]:
            cert_para = doc.add_paragraph(style="List Bullet")

            name_run = cert_para.add_run(cert.get("name", ""))
            name_run.bold = True

            if cert.get("issuer"):
                cert_para.add_run(f" - {cert['issuer']}")

            if cert.get("date"):
                cert_para.add_run(f" ({format_date(cert['date'])})")

        doc.add_paragraph()

    # Projects
    if content.get("projects"):
        add_heading("PROJECTS")

        for proj in content["projects"]:
            proj_title = doc.add_paragraph()
            proj_run = proj_title.add_run(proj.get("name", ""))
            proj_run.bold = True

            if proj.get("description"):
                desc_para = doc.add_paragraph()
                desc_para.add_run(proj["description"])

            if proj.get("technologies"):
                tech_para = doc.add_paragraph()
                tech_run = tech_para.add_run(f"Technologies: {', '.join(proj['technologies'])}")
                tech_run.italic = True

            links = []
            if proj.get("url"):
                links.append(f"Demo: {proj['url']}")
            if proj.get("github"):
                links.append(f"GitHub: {proj['github']}")

            if links:
                links_para = doc.add_paragraph()
                links_para.add_run(" | ".join(links))

            doc.add_paragraph()

    doc.save(str(filepath))
    return filepath


# ============== PDF GENERATION ==============

def generate_pdf(content: dict = None, filename: str = None) -> Path:
    """Generate a PDF resume file.

    Requires reportlab package.

    Args:
        content: Resume content dict (generated if not provided)
        filename: Output filename (auto-generated if not provided)

    Returns:
        Path to the generated file
    """
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
        from reportlab.lib.colors import HexColor
    except ImportError:
        raise ImportError(
            "reportlab is required for PDF generation. "
            "Install with: pip install reportlab"
        )

    if content is None:
        content = generate_resume_content()

    ensure_output_dir()

    if filename is None:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"resume_{timestamp}.pdf"

    filepath = OUTPUT_DIR / filename

    doc = SimpleDocTemplate(
        str(filepath),
        pagesize=letter,
        topMargin=0.5 * inch,
        bottomMargin=0.5 * inch,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
    )

    # Styles
    styles = getSampleStyleSheet()

    # Custom styles
    name_style = ParagraphStyle(
        "NameStyle",
        parent=styles["Heading1"],
        fontSize=18,
        textColor=HexColor("#1F497D"),
        spaceAfter=6,
        alignment=1,  # Center
    )

    title_style = ParagraphStyle(
        "TitleStyle",
        parent=styles["Normal"],
        fontSize=12,
        textColor=HexColor("#595959"),
        spaceAfter=6,
        alignment=1,
    )

    contact_style = ParagraphStyle(
        "ContactStyle",
        parent=styles["Normal"],
        fontSize=10,
        textColor=HexColor("#595959"),
        spaceAfter=12,
        alignment=1,
    )

    section_style = ParagraphStyle(
        "SectionStyle",
        parent=styles["Heading2"],
        fontSize=12,
        textColor=HexColor("#1F497D"),
        spaceBefore=12,
        spaceAfter=6,
        borderPadding=0,
    )

    heading_style = ParagraphStyle(
        "HeadingStyle",
        parent=styles["Normal"],
        fontSize=11,
        textColor=HexColor("#1F497D"),
        spaceAfter=2,
    )

    normal_style = ParagraphStyle(
        "NormalStyle",
        parent=styles["Normal"],
        fontSize=10,
        spaceAfter=4,
    )

    story = []

    # Header
    header = content["header"]

    story.append(Paragraph(header["name"], name_style))

    if header.get("title"):
        story.append(Paragraph(header["title"], title_style))

    contact_parts = []
    if header.get("email"):
        contact_parts.append(header["email"])
    if header.get("phone"):
        contact_parts.append(header["phone"])
    if header.get("location"):
        contact_parts.append(header["location"])

    if contact_parts:
        story.append(Paragraph(" | ".join(contact_parts), contact_style))

    social_parts = []
    if header.get("linkedin"):
        social_parts.append(f"LinkedIn: {header['linkedin']}")
    if header.get("github"):
        social_parts.append(f"GitHub: {header['github']}")
    if header.get("portfolio"):
        social_parts.append(f"Portfolio: {header['portfolio']}")

    if social_parts:
        story.append(Paragraph(" | ".join(social_parts), contact_style))

    # Summary
    if content.get("summary"):
        story.append(Paragraph("SUMMARY", section_style))
        story.append(Paragraph(content["summary"], normal_style))

    # Experience
    if content.get("experience"):
        story.append(Paragraph("EXPERIENCE", section_style))

        for exp in content["experience"]:
            title_line = f"<b>{exp.get('title', '')}</b> | {exp.get('company', '')}"
            story.append(Paragraph(title_line, normal_style))

            dates = f"{format_date(exp.get('start_date'))} - {format_date(exp.get('end_date'))}"
            if exp.get("location"):
                dates += f" | {exp['location']}"
            story.append(Paragraph(f"<i>{dates}</i>", normal_style))

            if exp.get("description"):
                story.append(Paragraph(exp["description"], normal_style))

            if exp.get("achievements"):
                for achievement in exp["achievements"]:
                    story.append(Paragraph(f"• {achievement}", normal_style))

            if exp.get("technologies"):
                story.append(Paragraph(f"<i>Technologies: {', '.join(exp['technologies'])}</i>", normal_style))

            story.append(Spacer(1, 6))

    # Education
    if content.get("education"):
        story.append(Paragraph("EDUCATION", section_style))

        for edu in content["education"]:
            degree_line = f"<b>{edu.get('degree', '')}"
            if edu.get("field"):
                degree_line += f" in {edu['field']}"
            degree_line += "</b>"
            story.append(Paragraph(degree_line, normal_style))

            story.append(Paragraph(edu.get("institution", ""), normal_style))

            if edu.get("gpa"):
                story.append(Paragraph(f"GPA: {edu['gpa']}", normal_style))

            story.append(Spacer(1, 6))

    # Skills
    if content.get("skills"):
        story.append(Paragraph("SKILLS", section_style))

        skills_by_cat = {}
        for skill in content["skills"]:
            cat = skill.get("category", "technical")
            if cat not in skills_by_cat:
                skills_by_cat[cat] = []
            skills_by_cat[cat].append(skill)

        for cat in ["technical", "soft", "language"]:
            if cat in skills_by_cat:
                cat_name = cat.title() + (" Skills" if cat != "language" else "s")
                skill_names = [s.get("name", "") for s in skills_by_cat[cat]]
                story.append(Paragraph(f"<b>{cat_name}:</b> {', '.join(skill_names)}", normal_style))

        story.append(Spacer(1, 6))

    # Certifications
    if content.get("certifications"):
        story.append(Paragraph("CERTIFICATIONS", section_style))

        for cert in content["certifications"]:
            cert_line = f"<b>{cert.get('name', '')}</b>"
            if cert.get("issuer"):
                cert_line += f" - {cert['issuer']}"
            if cert.get("date"):
                cert_line += f" ({format_date(cert['date'])})"
            story.append(Paragraph(f"• {cert_line}", normal_style))

        story.append(Spacer(1, 6))

    # Projects
    if content.get("projects"):
        story.append(Paragraph("PROJECTS", section_style))

        for proj in content["projects"]:
            story.append(Paragraph(f"<b>{proj.get('name', '')}</b>", normal_style))

            if proj.get("description"):
                story.append(Paragraph(proj["description"], normal_style))

            if proj.get("technologies"):
                story.append(Paragraph(f"<i>Technologies: {', '.join(proj['technologies'])}</i>", normal_style))

            story.append(Spacer(1, 6))

    doc.build(story)
    return filepath


# ============== GENERATE ALL ==============

def generate_all_formats(
    formats: list[Literal["pdf", "docx", "md"]] = None,
    filename_base: str = None,
) -> dict:
    """Generate resume in multiple formats.

    Args:
        formats: List of formats to generate (default: all)
        filename_base: Base filename without extension (auto-generated if not provided)

    Returns:
        Dict with paths to generated files
    """
    if formats is None:
        formats = ["md", "docx", "pdf"]

    if filename_base is None:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename_base = f"resume_{timestamp}"

    content = generate_resume_content()
    results = {}
    errors = {}

    # Always generate markdown first (no dependencies)
    if "md" in formats:
        try:
            results["md"] = str(generate_markdown(content, f"{filename_base}.md"))
        except Exception as e:
            errors["md"] = str(e)

    if "docx" in formats:
        try:
            results["docx"] = str(generate_docx(content, f"{filename_base}.docx"))
        except ImportError as e:
            errors["docx"] = "python-docx not installed. Run: pip install python-docx"
        except Exception as e:
            errors["docx"] = str(e)

    if "pdf" in formats:
        try:
            results["pdf"] = str(generate_pdf(content, f"{filename_base}.pdf"))
        except ImportError as e:
            errors["pdf"] = "reportlab not installed. Run: pip install reportlab"
        except Exception as e:
            errors["pdf"] = str(e)

    return {
        "files": results,
        "errors": errors,
        "content": content,
        "output_directory": str(OUTPUT_DIR),
    }


def list_generated_files() -> list:
    """List all generated resume files.

    Returns:
        List of file info dicts
    """
    ensure_output_dir()
    files = []

    for filepath in OUTPUT_DIR.iterdir():
        if filepath.is_file() and filepath.suffix in [".md", ".docx", ".pdf"]:
            stat = filepath.stat()
            files.append({
                "filename": filepath.name,
                "path": str(filepath),
                "size": stat.st_size,
                "created": datetime.fromtimestamp(stat.st_ctime).isoformat(),
                "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(),
            })

    return sorted(files, key=lambda x: x["modified"], reverse=True)


def delete_generated_file(filename: str) -> bool:
    """Delete a generated file.

    Args:
        filename: Name of file to delete

    Returns:
        True if deleted, False if not found
    """
    filepath = OUTPUT_DIR / filename
    if filepath.exists() and filepath.is_file():
        filepath.unlink()
        return True
    return False